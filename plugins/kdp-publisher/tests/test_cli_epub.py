import struct
import zlib

import docx
import ebooklib
from ebooklib import epub

from kdp_publisher.cli import main


def _png(w=4, h=4) -> bytes:
    raw = b"".join(b"\x00" + bytes((10, 20, 30)) * w for _ in range(h))

    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _doc(path):
    d = docx.Document()
    d.add_paragraph("Title: Epub Book")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("Body of chapter one.")
    d.add_heading("Chapter Two", level=1)
    d.add_paragraph("Body of chapter two.")
    d.save(str(path))


def _read(path):
    return epub.read_epub(str(path))


def test_cli_epub_writes_valid_epub(tmp_path):
    src = tmp_path / "in.docx"
    _doc(src)
    out = tmp_path / "book.epub"
    rc = main(["epub", str(src), "-o", str(out)])
    assert rc == 0
    assert out.read_bytes()[:2] == b"PK"
    b = _read(out)
    assert b.get_metadata("DC", "title")[0][0] == "Epub Book"


def test_cli_epub_with_cover(tmp_path):
    src = tmp_path / "in.docx"
    _doc(src)
    cover = tmp_path / "cover.png"
    cover.write_bytes(_png(8, 12))
    out = tmp_path / "book.epub"
    rc = main(["epub", str(src), "--cover-image", str(cover), "-o", str(out)])
    assert rc == 0
    b = _read(out)
    assert list(b.get_items_of_type(ebooklib.ITEM_COVER))
    spine_ids = [idref for idref, _linear in b.spine]
    assert spine_ids[0] == "cover"


def test_cli_epub_missing_metadata_errors(tmp_path):
    d = docx.Document()
    d.add_paragraph("Title: Only Title")
    d.add_heading("Ch", level=1)
    d.add_paragraph("body")
    src = tmp_path / "bad.docx"
    d.save(str(src))
    rc = main(["epub", str(src), "-o", str(tmp_path / "o.epub")])
    assert rc != 0
