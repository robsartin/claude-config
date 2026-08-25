import docx
import pypdf

from kdp_publisher.cli import main


def _doc(path):
    d = docx.Document()
    d.add_paragraph("Title: CLI Book")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    for i in range(600):
        d.add_paragraph(f"Paragraph {i} with enough words to fill a page eventually.")
    d.save(path)


def test_cli_renders_interior(tmp_path):
    src = tmp_path / "in.docx"
    _doc(str(src))
    out = tmp_path / "out.pdf"
    rc = main(["interior", str(src), "-o", str(out)])
    assert rc == 0
    assert out.exists() and out.stat().st_size > 1000


def test_cli_rejects_under_min_pages(tmp_path, capsys):
    d = docx.Document()
    d.add_paragraph("Title: Too Short")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    for i in range(5):
        d.add_paragraph(f"Paragraph {i}.")
    src = tmp_path / "short.docx"
    d.save(str(src))
    out = tmp_path / "out.pdf"

    rc = main(["interior", str(src), "-o", str(out)])

    captured = capsys.readouterr().out
    assert rc != 0
    assert "24" in captured
    assert not out.exists()


def test_cli_errors_on_missing_metadata(tmp_path):
    d = docx.Document()
    d.add_paragraph("Title: Only Title")
    d.add_heading("Ch", level=1)
    d.add_paragraph("body")
    src = tmp_path / "bad.docx"
    d.save(str(src))
    rc = main(["interior", str(src), "-o", str(tmp_path / "o.pdf")])
    assert rc != 0


def test_cli_missing_metadata_message_suggests_correct_paper_flag(tmp_path, capsys):
    d = docx.Document()
    d.add_paragraph("Title: Only Title")
    d.add_heading("Ch", level=1)
    d.add_paragraph("body")
    src = tmp_path / "bad.docx"
    d.save(str(src))
    rc = main(["interior", str(src), "-o", str(tmp_path / "o.pdf")])
    out = capsys.readouterr().out
    assert rc != 0
    assert "--paper" in out
    assert "--paper_type" not in out


def test_cli_invalid_trim_override_is_rejected_cleanly(tmp_path, capsys):
    d = docx.Document()
    d.add_heading("Ch", level=1)
    d.add_paragraph("body")
    src = tmp_path / "bad_trim.docx"
    d.save(str(src))
    rc = main(
        [
            "interior",
            str(src),
            "-o",
            str(tmp_path / "o.pdf"),
            "--title",
            "T",
            "--author",
            "A",
            "--trim",
            "7x10",
            "--paper",
            "cream",
        ]
    )
    out = capsys.readouterr().out
    assert isinstance(rc, int)
    assert rc != 0
    assert "6x9" in out or "5x8" in out or "8.5x11" in out


def test_cli_normalizes_case_and_alias_overrides(tmp_path):
    d = docx.Document()
    d.add_heading("Ch", level=1)
    for i in range(600):
        d.add_paragraph(f"Paragraph {i} with enough words to fill a page eventually.")
    src = tmp_path / "ok.docx"
    d.save(str(src))
    out = tmp_path / "out.pdf"
    rc = main(
        [
            "interior",
            str(src),
            "-o",
            str(out),
            "--title",
            "T",
            "--author",
            "A",
            "--trim",
            "6X9",
            "--paper",
            "CREAM",
        ]
    )
    assert rc == 0
    assert out.exists()


def _short_doc(path, trim="6x9"):
    d = docx.Document()
    d.add_paragraph("Title: Too Short")
    d.add_paragraph("Author: Tester")
    d.add_paragraph(f"Trim: {trim}")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    for i in range(5):
        d.add_paragraph(f"Paragraph {i}.")
    d.save(str(path))


def _google_pdf(path, width_in, height_in, pages):
    w = pypdf.PdfWriter()
    for _ in range(pages):
        w.add_blank_page(width=width_in * 72, height=height_in * 72)
    with open(path, "wb") as f:
        w.write(f)


def test_short_interior_still_prints_the_validation_report(tmp_path, capsys):
    """A draft is short exactly while it is being set up — the window where a
    wrong Page setup is cheapest to fix. Failing on page count must not hide it."""
    src = tmp_path / "short.docx"
    _short_doc(src)
    google = tmp_path / "google.pdf"
    _google_pdf(google, 5.0, 8.0, 6)  # 5x8 export against a 6x9 manuscript

    rc = main(["interior", str(src), "--google-pdf", str(google), "-o", str(tmp_path / "o.pdf")])

    out = capsys.readouterr().out
    assert rc != 0
    assert "trim" in out
    assert "fonts" in out
    # the page-count error stays the last thing on screen
    assert out.rstrip().splitlines()[-1].startswith("error:")


def test_short_cover_still_prints_the_validation_report(tmp_path, capsys):
    src = tmp_path / "short.docx"
    _short_doc(src)
    google = tmp_path / "google.pdf"
    _google_pdf(google, 5.0, 8.0, 6)

    rc = main(["cover", str(src), "--google-pdf", str(google), "-o", str(tmp_path / "spec.txt")])

    out = capsys.readouterr().out
    assert rc != 0
    assert "trim" in out
    assert out.rstrip().splitlines()[-1].startswith("error:")
