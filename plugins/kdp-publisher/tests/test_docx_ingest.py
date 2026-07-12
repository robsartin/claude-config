import docx

from kdp_publisher.ingest.docx_ingest import ingest_docx


def _make_docx(path):
    d = docx.Document()
    d.add_paragraph("Title: Fixture Book")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("First paragraph.")
    d.add_heading("A Section", level=2)
    d.add_paragraph("Second paragraph.")
    d.add_heading("Chapter Two", level=1)
    d.add_paragraph("Another paragraph.")
    d.save(path)


def test_ingest_builds_book_from_docx(tmp_path):
    p = tmp_path / "book.docx"
    _make_docx(p)
    book, missing = ingest_docx(str(p))
    assert missing == []
    assert book.metadata.title == "Fixture Book"
    assert book.metadata.paper_type == "cream-bw"
    assert [c.title for c in book.chapters] == ["Chapter One", "Chapter Two"]
    # Chapter One has: H2 heading + 2 paragraphs
    titles = book.chapters[0]
    assert any(getattr(b, "text", "") == "First paragraph." for b in titles.blocks)


def test_overrides_fill_missing_fields(tmp_path):
    d = docx.Document()
    d.add_paragraph("Title: No Paper Book")
    d.add_heading("Ch", level=1)
    d.add_paragraph("body")
    p = tmp_path / "b2.docx"
    d.save(str(p))
    book, missing = ingest_docx(
        str(p), overrides={"author": "X", "trim": "6x9", "paper_type": "white-bw"}
    )
    assert missing == []
    assert book.metadata.author == "X"
    assert book.metadata.paper_type == "white-bw"
