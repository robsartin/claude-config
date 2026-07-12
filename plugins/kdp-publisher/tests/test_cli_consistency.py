import re
import struct
import zlib

import docx
import pytest

from kdp_publisher import kdp_rules as r
from kdp_publisher.cli import main
from kdp_publisher.ingest.docx_ingest import ingest_docx
from kdp_publisher.interior.pipeline import produce_interior


def _png(w=40, h=60, rgb=(200, 120, 60)) -> bytes:
    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))

    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _doc(path, paras=600):
    d = docx.Document()
    d.add_paragraph("Title: Consistency Book")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    for i in range(paras):
        d.add_paragraph(f"Paragraph {i} with several words to fill the page.")
    d.save(str(path))


def test_interior_and_cover_report_same_page_count(tmp_path, capsys):
    """`interior` and `cover` on the same docx must agree on the page count --
    the cover's spine width is derived from it, so a mismatch would silently
    produce a wraparound cover that doesn't match the interior the author
    uploads to KDP.
    """
    src = tmp_path / "in.docx"
    _doc(src)

    # Ground truth: what the interior pipeline actually produces for this
    # docx, computed directly (independent of any CLI print statement).
    book, missing = ingest_docx(str(src))
    assert not missing
    true_page_count = produce_interior(book, None).page_count

    interior_out = tmp_path / "interior.pdf"
    rc = main(["interior", str(src), "-o", str(interior_out)])
    assert rc == 0
    interior_stdout = capsys.readouterr().out
    interior_match = re.search(r"interior: (\d+) pages", interior_stdout)
    assert interior_match, f"could not find page count in: {interior_stdout!r}"
    interior_pages = int(interior_match.group(1))
    assert interior_pages == true_page_count

    img = tmp_path / "front.png"
    img.write_bytes(_png())
    cover_out = tmp_path / "cover.pdf"
    rc = main(["cover", str(src), "--cover-image", str(img), "-o", str(cover_out)])
    assert rc == 0
    cover_stdout = capsys.readouterr().out
    cover_match = re.search(r"cover: .* (\d+) pages, spine ([\d.]+) in", cover_stdout)
    assert cover_match, f"could not find page count/spine in: {cover_stdout!r}"
    cover_pages = int(cover_match.group(1))
    cover_spine_in = float(cover_match.group(2))

    assert cover_pages == interior_pages == true_page_count

    # The spine width is a function of page count; this is what actually
    # catches a divergence if the cover subcommand ever computed its
    # CoverSpec from a different page count than the one it prints/than the
    # interior pipeline reports -- printing `result.page_count` alone
    # wouldn't reveal that, since the spec's own page_count is a separate
    # value baked into spine_in/width_in.
    assert cover_spine_in == pytest.approx(r.spine_in(true_page_count, "cream-bw"), abs=1e-3)
