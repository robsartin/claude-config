import io
import struct
import zlib

import docx
import pypdf

from kdp_publisher.cli import main


def _png(w=40, h=60, rgb=(200, 120, 60)) -> bytes:
    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))

    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def _doc(path, paras=600):
    d = docx.Document()
    d.add_paragraph("Title: Cover Book")
    d.add_paragraph("Author: Tester")
    d.add_paragraph("Trim: 6x9")
    d.add_paragraph("Paper: cream")
    d.add_heading("Chapter One", level=1)
    for i in range(paras):
        d.add_paragraph(f"Paragraph {i} with several words to fill the page.")
    d.save(str(path))


def test_cover_with_image_writes_pdf(tmp_path):
    src = tmp_path / "in.docx"
    _doc(src)
    img = tmp_path / "front.png"
    img.write_bytes(_png())
    out = tmp_path / "cover.pdf"
    rc = main(["cover", str(src), "--cover-image", str(img), "-o", str(out)])
    assert rc == 0
    assert len(pypdf.PdfReader(io.BytesIO(out.read_bytes())).pages) == 1


def test_cover_without_image_writes_spec_and_prompt(tmp_path):
    src = tmp_path / "in.docx"
    _doc(src)
    out = tmp_path / "cover-spec.txt"
    rc = main(["cover", str(src), "-o", str(out)])
    assert rc == 0
    text = out.read_text()
    assert "px at 300 DPI" in text
    assert "WRAPAROUND" in text
    assert "spine" in text.lower()
